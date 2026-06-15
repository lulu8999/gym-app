#!/usr/bin/env python3
"""
Word 文档工具函数集
提供文档分析、修复、格式化等功能
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional, Tuple
import re


def analyze_docx_structure(docx_path: str) -> dict:
    """
    分析 Word 文档结构，返回详细信息
    """
    doc = Document(docx_path)

    # 统计信息
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'empty_paragraphs': 0,
        'short_paragraphs': 0,  # < 20 字
        'sections': [],
        'table_like_regions': []
    }

    # 识别章节标题
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            stats['empty_paragraphs'] += 1
        elif len(text) < 20:
            stats['short_paragraphs'] += 1

        # 识别章节标题
        if text.startswith('第') and ('章' in text or '节' in text):
            stats['sections'].append({
                'index': i,
                'title': text[:50],
                'level': 1 if '章' in text else 2
            })

    return stats


def find_table_like_regions(docx_path: str, min_rows: int = 3) -> List[Tuple[int, int]]:
    """
    找出可能是表格被拆散的区域

    Args:
        docx_path: Word 文件路径
        min_rows: 最小行数（认为是表格的阈值）

    Returns:
        列表，每项为 (start_idx, end_idx) 表示一个区域
    """
    doc = Document(docx_path)
    regions = []
    current_region = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # 如果是空行，结束当前区域
        if not text:
            if len(current_region) >= min_rows:
                regions.append((current_region[0][0], current_region[-1][0]))
            current_region = []
            continue

        # 判断是否是可能的表格内容（短文本）
        if len(text) < 30:
            current_region.append((i, text))
        else:
            if len(current_region) >= min_rows:
                regions.append((current_region[0][0], current_region[-1][0]))
            current_region = []

    # 处理文档末尾的区域
    if len(current_region) >= min_rows:
        regions.append((current_region[0][0], current_region[-1][0]))

    return regions


def clean_empty_paragraphs(docx_path: str, output_path: Optional[str] = None) -> str:
    """
    清理空段落

    Args:
        docx_path: 输入文件路径
        output_path: 输出文件路径（默认覆盖原文件）

    Returns:
        输出文件路径
    """
    doc = Document(docx_path)

    # 收集需要保留的段落
    new_paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # 保留非空段落
            new_paragraphs.append(text)

    # 清空文档
    for para in doc.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)

    # 重新添加
    for text in new_paragraphs:
        doc.add_paragraph(text)

    output = output_path or docx_path
    doc.save(output)
    return output


def extract_tables_to_markdown(docx_path: str) -> List[str]:
    """
    提取文档中的表格为 Markdown 格式

    Returns:
        每个表格的 Markdown 表示列表
    """
    doc = Document(docx_path)
    md_tables = []

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')

        if len(rows) >= 2:
            # 添加分隔行
            separator = '|' + '|'.join(['---'] * len(table.rows[0].cells)) + '|'
            rows.insert(1, separator)
            md_tables.append('\n'.join(rows))

    return md_tables


def set_paragraph_format(
    docx_path: str,
    output_path: str,
    font_name: str = '宋体',
    font_size: int = 12,
    line_spacing: float = 1.5,
    first_line_indent: float = 0.74  # 2字符约 0.74cm
):
    """
    设置段落格式（公文排版风格）
    """
    doc = Document(docx_path)

    for para in doc.paragraphs:
        # 设置段落格式
        pf = para.paragraph_format
        pf.line_spacing = line_spacing
        pf.first_line_indent = Cm(first_line_indent)

        # 设置字体
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("用法: python docx_utils.py <command> <args>")
        print("")
        print("命令:")
        print("  analyze <docx>        - 分析文档结构")
        print("  clean <docx> [out]    - 清理空段落")
        print("  extract <docx>        - 提取表格为 Markdown")
        sys.exit(1)

    cmd = sys.argv[1]

    if cmd == 'analyze' and len(sys.argv) >= 3:
        result = analyze_docx_structure(sys.argv[2])
        print(f"文档分析结果:")
        print(f"  总段落: {result['total_paragraphs']}")
        print(f"  空段落: {result['empty_paragraphs']}")
        print(f"  短段落: {result['short_paragraphs']}")
        print(f"  表格数: {result['total_tables']}")
        print(f"  章节数: {len(result['sections'])}")

    elif cmd == 'clean' and len(sys.argv) >= 3:
        out = sys.argv[3] if len(sys.argv) > 3 else None
        result = clean_empty_paragraphs(sys.argv[2], out)
        print(f"清理完成: {result}")

    elif cmd == 'extract' and len(sys.argv) >= 3:
        tables = extract_tables_to_markdown(sys.argv[2])
        for i, table in enumerate(tables):
            print(f"\n表格 {i+1}:")
            print(table)
