#!/usr/bin/env python3
"""
Markdown → Word 自动转换工具
解决 AI 生成 Word 文档时表格被拆散的问题

使用流程：
1. AI 生成标准 Markdown 格式内容
2. 本工具自动调用 Pandoc 转换为 Word
3. 生成格式正确的 .docx 文件
"""

import subprocess
import tempfile
import os
import sys
from pathlib import Path
from typing import Optional, List


def md_to_docx(
    content: str,
    output_path: str,
    title: Optional[str] = None,
    toc: bool = True,
    extra_args: Optional[List[str]] = None
) -> str:
    """
    将 Markdown 内容转换为 Word 文档

    Args:
        content: Markdown 格式的内容
        output_path: 输出的 Word 文件路径
        title: 文档标题（可选）
        toc: 是否生成目录
        extra_args: 额外的 pandoc 参数

    Returns:
        输出文件的绝对路径

    Raises:
        RuntimeError: 转换失败
    """
    # 确保输出目录存在
    output_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(output_dir, exist_ok=True)

    # 添加 YAML frontmatter（如果有标题）
    if title:
        full_content = f"---\ntitle: {title}\n---\n\n{content}"
    else:
        full_content = content

    # 创建临时 markdown 文件
    with tempfile.NamedTemporaryFile(
        mode='w',
        suffix='.md',
        delete=False,
        encoding='utf-8'
    ) as f:
        f.write(full_content)
        md_path = f.name

    try:
        # 构建 pandoc 命令
        cmd = ['pandoc', md_path, '-o', output_path]

        if toc:
            cmd.append('--toc')

        if extra_args:
            cmd.extend(extra_args)

        # 执行转换
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True
        )

        if result.returncode != 0:
            raise RuntimeError(f"Pandoc 转换失败: {result.stderr}")

        return os.path.abspath(output_path)

    finally:
        # 清理临时文件
        os.unlink(md_path)


def docx_to_md(input_path: str) -> str:
    """
    将 Word 文档转换为 Markdown（用于分析现有文档）

    Args:
        input_path: Word 文件路径

    Returns:
        Markdown 格式的内容
    """
    with tempfile.NamedTemporaryFile(
        mode='w',
        suffix='.md',
        delete=False,
        encoding='utf-8'
    ) as f:
        md_path = f.name

    try:
        result = subprocess.run(
            ['pandoc', input_path, '-o', md_path],
            capture_output=True,
            text=True
        )

        if result.returncode != 0:
            raise RuntimeError(f"Pandoc 转换失败: {result.stderr}")

        with open(md_path, 'r', encoding='utf-8') as f:
            return f.read()

    finally:
        os.unlink(md_path)


def validate_docx(docx_path: str) -> dict:
    """
    验证 Word 文档质量

    Returns:
        包含段落数、表格数、可能问题的字典
    """
    try:
        from docx import Document
        doc = Document(docx_path)

        # 检查空行
        empty_count = sum(1 for p in doc.paragraphs if not p.text.strip())

        # 检查短段落（可能是表格被拆散）
        short_count = sum(1 for p in doc.paragraphs if 0 < len(p.text.strip()) < 20)

        return {
            'valid': True,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'empty_lines': empty_count,
            'short_paragraphs': short_count,
            'issues': []
        }

    except Exception as e:
        return {
            'valid': False,
            'error': str(e),
            'issues': [str(e)]
        }


if __name__ == '__main__':
    # CLI 接口
    if len(sys.argv) < 3:
        print("用法: python md_to_docx.py <input.md> <output.docx> [--toc]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]
    toc = '--toc' in sys.argv

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    result = md_to_docx(content, output_file, toc=toc)
    print(f"✓ 转换成功: {result}")

    # 验证结果
    validation = validate_docx(result)
    print(f"  段落数: {validation['paragraphs']}")
    print(f"  表格数: {validation['tables']}")
    print(f"  空行数: {validation['empty_lines']}")
